# -*- coding: utf-8 -*-
"""
Insert new sections into baocao.docx:
  - 4.2.3. Danh gia Document RAG Pipeline
  - 1.4.3. BM25
  - 1.4.4. Hybrid Search + RRF
  - 1.5.1. Bien doi truy van (HyDE, Multi-Query, Decomposition)
  - 1.5.2. Reranking
  - 1.5.3. OOS Detection
  - 1.5.4. Optuna TPE
  - 3.3.3. Indexing Document RAG
  - 3.3.4. Retrieval Pipeline
  - 3.3.5. Safety & UX
  - 3.3.6. Toi uu tham so Optuna
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os

doc = Document("baocao.docx")

# ─── Helpers ───────────────────────────────────────────────────
def find_para_index(doc, text_contains):
    for i, p in enumerate(doc.paragraphs):
        if text_contains.lower() in p.text.lower():
            return i
    return -1

def insert_para_after(doc, ref_para_idx, text, style_name, bold=False, size=None):
    """Insert a new paragraph after ref_para_idx. Returns new para index."""
    ref_para = doc.paragraphs[ref_para_idx]
    new_para = OxmlElement("w:p")
    ref_para._element.addnext(new_para)
    # re-find inserted para (it is now ref_para_idx+1)
    para = doc.paragraphs[ref_para_idx + 1]
    para.style = doc.styles[style_name]
    run = para.add_run(text)
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    return ref_para_idx + 1

def insert_text_after(doc, ref_para_idx, text, style_name="Normal", indent_cm=0):
    ref_para = doc.paragraphs[ref_para_idx]
    new_para = OxmlElement("w:p")
    ref_para._element.addnext(new_para)
    para = doc.paragraphs[ref_para_idx + 1]
    para.style = doc.styles[style_name]
    para.add_run(text)
    if indent_cm:
        pf = para.paragraph_format
        pf.left_indent = Cm(indent_cm)
    return ref_para_idx + 1

def insert_bullet_after(doc, ref_para_idx, text):
    ref_para = doc.paragraphs[ref_para_idx]
    new_para = OxmlElement("w:p")
    ref_para._element.addnext(new_para)
    para = doc.paragraphs[ref_para_idx + 1]
    try:
        para.style = doc.styles["List Bullet"]
    except:
        para.style = doc.styles["Normal"]
        para.paragraph_format.left_indent = Cm(1.0)
    para.add_run(text)
    return ref_para_idx + 1

def insert_image_after(doc, ref_para_idx, img_path, caption, width_cm=15):
    ref_para = doc.paragraphs[ref_para_idx]
    # Image para
    new_para = OxmlElement("w:p")
    ref_para._element.addnext(new_para)
    img_para = doc.paragraphs[ref_para_idx + 1]
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(img_path):
        run = img_para.add_run()
        run.add_picture(img_path, width=Cm(width_cm))
    else:
        img_para.add_run(f"[Hình: {img_path} — không tìm thấy]")
    # Caption para
    new_para2 = OxmlElement("w:p")
    img_para._element.addnext(new_para2)
    cap_para = doc.paragraphs[ref_para_idx + 2]
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = cap_para.add_run(caption)
    run2.italic = True
    run2.font.size = Pt(10)
    return ref_para_idx + 2

def insert_table_after(doc, ref_para_idx, headers, rows, caption):
    """Insert a table after ref_para_idx. Returns new index after table."""
    ref_para = doc.paragraphs[ref_para_idx]
    # Insert table via body
    body = doc.element.body
    ref_elem = ref_para._element
    tbl = doc.add_table(rows=len(rows)+1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # shade header
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F3864")
        shd.set(qn("w:color"), "FFFFFF")
        tcPr.append(shd)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.rows[i+1].cells[j]
            cell.text = str(val)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i % 2 == 1:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "EBF3FB")
                tcPr.append(shd)
    # Move table to correct position
    tbl_elem = tbl._tbl
    body.remove(tbl_elem)
    ref_elem.addnext(tbl_elem)
    # Caption after table
    cap_elem = OxmlElement("w:p")
    tbl_elem.addnext(cap_elem)
    new_idx = ref_para_idx + 2  # approximate; table = 1 elem, caption = 1 para
    para = doc.paragraphs[ref_para_idx + 2] if ref_para_idx + 2 < len(doc.paragraphs) else doc.paragraphs[-1]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(caption)
    run.italic = True
    run.font.size = Pt(10)
    return ref_para_idx + 2

# ══════════════════════════════════════════════════════════════
# SECTION 4.2.3 — Insert after "4.2.2. Đánh giá hiệu năng"
# ══════════════════════════════════════════════════════════════
idx = find_para_index(doc, "4.2.2")
if idx < 0:
    idx = find_para_index(doc, "Đánh giá hiệu năng và độ trễ")
print(f"Found 4.2.2 at para index: {idx}")

# Find end of 4.2.2 section (before 4.3 or KET LUAN)
end_422 = idx
for i in range(idx+1, len(doc.paragraphs)):
    p = doc.paragraphs[i]
    txt = p.text.strip()
    if p.style.name.startswith("Heading") and ("4.3" in txt or "KẾT LUẬN" in txt or "Giao diện" in txt):
        break
    end_422 = i

cur = end_422

# --- Heading 4.2.3 ---
cur = insert_para_after(doc, cur, "4.2.3. Đánh giá hệ thống hỏi đáp tài liệu (Document RAG)", "Heading 3")

# --- 4.2.3.1 ---
cur = insert_para_after(doc, cur, "4.2.3.1. Thiết kế bộ dữ liệu đánh giá", "Heading 4")
cur = insert_text_after(doc, cur,
    "Để đánh giá toàn diện chất lượng hệ thống Document RAG, nhóm xây dựng bộ câu hỏi đánh giá "
    "gồm 129 câu hỏi phân chia thành ba tập độc lập. Tập Dev Set (72 câu) được sử dụng trong quá "
    "trình tối ưu tham số Optuna, không được dùng để đánh giá cuối cùng. Tập Holdout Set (32 câu) "
    "được giữ kín trong suốt quá trình phát triển nhằm đo lường khả năng tổng quát hóa. Tập Final "
    "Set (25 câu) là bộ đánh giá chính thức cuối cùng, chỉ được chạy một lần duy nhất khi hệ thống "
    "đã ổn định.")
cur = insert_text_after(doc, cur,
    "Mỗi câu hỏi được phân loại theo 5 category phản ánh các dạng truy vấn thực tế: fact (câu hỏi "
    "tra cứu số liệu/sự kiện cụ thể), comparison (so sánh giữa hai hoặc nhiều thực thể), logic "
    "(suy luận nhiều bước từ nhiều đoạn tài liệu), multi-file (câu hỏi đòi hỏi tổng hợp từ nhiều "
    "tài liệu khác nhau), và out-of-scope (câu hỏi nằm ngoài phạm vi tài liệu được index — hệ "
    "thống phải từ chối trả lời).")
cur = insert_image_after(doc, cur,
    "chart_dataset_balance.png",
    "Hình 4.X. Phân bố loại câu hỏi trong 3 bộ đánh giá — đảm bảo tính cân bằng và đại diện", 14)

# --- 4.2.3.2 ---
cur = insert_para_after(doc, cur, "4.2.3.2. Phương pháp chấm điểm tự động", "Heading 4")
cur = insert_text_after(doc, cur,
    "Hệ thống sử dụng LLM-as-Judge với mô hình qwen2.5:7b chạy trên Ollama để chấm điểm tự động. "
    "Mỗi câu trả lời được đánh giá theo hai tiêu chí: (1) Pass/Fail — đạt hay không đạt dựa trên "
    "việc trả lời đúng các điểm chính (expected_points), và (2) Point Recall — tỉ lệ số điểm "
    "được trả lời đúng trên tổng số điểm yêu cầu. Đối với câu hỏi out-of-scope, tiêu chí là "
    "Refuse Accuracy — hệ thống có từ chối đúng hay không. Mỗi câu được chạy ở chế độ mode=high "
    "(kích hoạt toàn bộ pipeline: HyDE, multi-query, decomposition, LLM reranking).")

# --- 4.2.3.3 ---
cur = insert_para_after(doc, cur, "4.2.3.3. Kết quả tổng thể", "Heading 4")
cur = insert_text_after(doc, cur,
    "Kết quả đánh giá trên ba bộ dữ liệu được tổng hợp trong Bảng 4.X. Hệ thống đạt 96.0% độ "
    "chính xác trên Final Set — bộ đánh giá chính thức cuối cùng — cho thấy hiệu quả cao trong "
    "điều kiện thực tế. Độ chính xác trên Holdout Set (84.4%) thấp hơn Dev Set (87.5%) là kết "
    "quả bình thường do Holdout được thiết kế để kiểm tra khả năng tổng quát hóa trên dữ liệu "
    "chưa từng thấy trong quá trình tối ưu.")

# Table tổng hợp
headers_tbl = ["Chỉ số", "Dev Set\n(72 câu)", "Holdout Set\n(32 câu)", "Final Set\n(25 câu)"]
rows_tbl = [
    ["Số câu Pass",          "63 / 72",   "27 / 32",   "24 / 25"],
    ["Độ chính xác (%)",     "87.5%",     "84.4%",     "96.0%"],
    ["Avg Recall (in-domain)","80.1%",    "72.9%",     "84.9%"],
    ["Refuse Accuracy (OOS)","72.7%",     "100.0%",    "100.0%"],
    ["Số câu OOS",           "11",        "8",         "4"],
    ["Pass-partial (recall<100%)","15",   "2",         "5"],
]
cur = insert_table_after(doc, cur, headers_tbl, rows_tbl,
    "Bảng 4.X. Tổng hợp kết quả đánh giá hệ thống Document RAG trên 3 bộ dữ liệu")

cur = insert_image_after(doc, cur, "chart1_overall_accuracy.png",
    "Hình 4.X+1. Độ chính xác tổng thể trên 3 bộ đánh giá", 13)

# --- 4.2.3.4 ---
cur = insert_para_after(doc, cur, "4.2.3.4. Phân tích theo loại câu hỏi", "Heading 4")
cur = insert_text_after(doc, cur,
    "Hình 4.X+2 trình bày độ chính xác phân theo 5 loại câu hỏi. Hệ thống đạt hiệu quả cao nhất "
    "ở hai loại multi-file (100% trên cả 3 bộ) và comparison (100% trên Final Set). Loại fact "
    "đạt 80% trên Final Set do một số câu hỏi yêu cầu số liệu cụ thể từ bảng OCR tiếng Nhật bị "
    "tokenize rời, khiến mô hình không map được số liệu đúng. Đặc biệt, out-of-scope đạt 100% "
    "refuse accuracy trên cả Holdout lẫn Final Set, chứng minh cơ chế OOS Detection hoạt động "
    "đáng tin cậy với ngưỡng threshold=0.42 được chỉnh thủ công sau khi Optuna đề xuất 0.539.")
cur = insert_image_after(doc, cur, "chart2_category_accuracy.png",
    "Hình 4.X+2. Độ chính xác theo loại câu hỏi trên 3 bộ đánh giá", 14)
cur = insert_image_after(doc, cur, "chart3_recall_distribution.png",
    "Hình 4.X+3. Phân bố Point Recall — mức độ đầy đủ của câu trả lời", 14)

# --- 4.2.3.5 ---
cur = insert_para_after(doc, cur, "4.2.3.5. Phân tích nguyên nhân lỗi", "Heading 4")
cur = insert_text_after(doc, cur,
    "Trong tổng số 9 câu sai trên Dev Set, phần lớn thuộc dạng Partial Recall — câu trả lời đúng "
    "một phần nhưng thiếu ý. Nguyên nhân chính được xác định gồm: (1) OCR text tiếng Nhật bị "
    "tokenize rời (đặc biệt với số liệu trong bảng), khiến BM25 và dense retrieval đều không "
    "map được chunk chứa số liệu đúng; (2) Câu hỏi multi-hop đòi hỏi tổng hợp từ nhiều trang "
    "xa nhau trong tài liệu, vượt quá giới hạn context window; (3) Một số câu hỏi out-of-scope "
    "bị hệ thống trả lời thay vì từ chối, do threshold OOS 0.42 khá thấp (ưu tiên ít false "
    "refusal hơn).")
cur = insert_image_after(doc, cur, "chart_error_analysis.png",
    "Hình 4.X+4. Phân tích nguyên nhân câu trả lời sai — Dev Set (72 câu)", 13)

# --- 4.2.3.6 ---
cur = insert_para_after(doc, cur, "4.2.3.6. Ảnh hưởng của tối ưu tham số Optuna", "Heading 4")
cur = insert_text_after(doc, cur,
    "Quá trình tối ưu tham số bằng Optuna TPE trên 50 trials (sweep2.db) cho thấy hai tham số "
    "có ảnh hưởng lớn nhất đến objective score là DecompositionSubQueryLimit "
    "(fANOVA importance = 0.241) và OutOfScopeScoreThreshold (0.198). "
    "Bộ tham số tối ưu (trial #22, obj=0.7321) cải thiện đáng kể so với giá trị mặc định ban đầu, "
    "đặc biệt ở khả năng recall cho câu hỏi phức tạp (logic và comparison). "
    "Tuy nhiên, ngưỡng OutOfScopeScoreThreshold do Optuna đề xuất (0.539) được "
    "điều chỉnh thủ công xuống 0.42 sau khi quan sát thực tế, nhằm ưu tiên giảm false refusal "
    "cho câu hỏi in-domain — đây là trade-off giữa OOS detection và in-domain recall được "
    "phân tích chi tiết trong Hình 4.X+5.")
cur = insert_image_after(doc, cur, "chart_bm25_heatmap.png",
    "Hình 4.X+5. Bề mặt objective theo BM25 k1 × b từ 50 Optuna trials", 14)
cur = insert_image_after(doc, cur, "chart_score_distribution.png",
    "Hình 4.X+6. Phân bố retrieval score — in-domain vs out-of-scope", 14)

# ══════════════════════════════════════════════════════════════
# SECTION 1.4.3 — BM25 — after "1.4.2. Chiến lược phân đoạn"
# ══════════════════════════════════════════════════════════════
idx_142 = find_para_index(doc, "1.4.2")
end_142 = idx_142
for i in range(idx_142+1, len(doc.paragraphs)):
    p = doc.paragraphs[i]
    if p.style.name.startswith("Heading") and p.text.strip() and "1.4.2" not in p.text:
        break
    end_142 = i

cur2 = end_142
cur2 = insert_para_after(doc, cur2, "1.4.3. Thuật toán tìm kiếm từ khóa BM25", "Heading 3")
cur2 = insert_text_after(doc, cur2,
    "BM25 (Best Matching 25) là thuật toán tìm kiếm từ khóa dựa trên thống kê, được coi là chuẩn "
    "vàng trong Information Retrieval cổ điển [Robertson et al., 1994]. Trong hệ thống Document "
    "RAG, BM25 đóng vai trò sparse retriever — bổ sung cho dense vector search ở những trường "
    "hợp truy vấn cần khớp chính xác từ khóa, tên riêng, mã hiệu kỹ thuật hoặc số liệu cụ thể "
    "mà dense embedding đôi khi bỏ sót.")
cur2 = insert_text_after(doc, cur2,
    "Điểm nổi bật của BM25 so với TF-IDF truyền thống là cơ chế bão hòa tần suất (TF Saturation): "
    "thay vì điểm số tăng tuyến tính vô hạn theo số lần xuất hiện từ khóa, BM25 giới hạn mức tăng "
    "tiệm cận về giá trị k1+1, ngăn chặn hiện tượng spam từ khóa. Ngoài ra, tham số b kiểm soát "
    "chuẩn hóa độ dài chunk — rất quan trọng với text OCR tiếng Nhật vốn có độ dài không đồng đều "
    "giữa các trang. Hệ thống tự cài đặt BM25 trong C# với tham số k1=1.34 và b=0.57 được xác "
    "định qua Optuna.")
cur2 = insert_image_after(doc, cur2, "concept_bm25.png",
    "Hình 3.X. BM25 — TF Saturation (trái) và Length Normalization (phải) với k1=1.34, b=0.57", 14)

# --- 1.4.4. Hybrid Search + RRF ---
cur2 = insert_para_after(doc, cur2, "1.4.4. Hybrid Search và Reciprocal Rank Fusion", "Heading 3")
cur2 = insert_text_after(doc, cur2,
    "Hybrid Search kết hợp song song hai luồng tìm kiếm: Dense Retrieval (vector semantic) và "
    "Sparse Retrieval (BM25 keyword). Mỗi luồng trả về danh sách kết quả theo thang điểm riêng "
    "không thể so sánh trực tiếp. Để hợp nhất hai danh sách này, hệ thống sử dụng Reciprocal Rank "
    "Fusion (RRF) [Cormack et al., 2009] — phương pháp chỉ dựa trên thứ hạng, không phụ thuộc "
    "vào thang điểm tuyệt đối.")
cur2 = insert_text_after(doc, cur2,
    "Công thức RRF: Score(d) = Σ 1/(k + rank_i(d)), trong đó k là hằng số làm mượt (k=44 theo "
    "Optuna), rank_i(d) là thứ hạng của chunk d trong danh sách i. Chunk nào xuất hiện ở thứ "
    "hạng cao trong cả hai luồng sẽ được ưu tiên — đây là cơ chế đồng thuận giúp lọc noise "
    "hiệu quả hơn so với từng phương pháp đơn lẻ.")
cur2 = insert_image_after(doc, cur2, "concept_rrf.png",
    "Hình 3.X+1. Ví dụ minh họa RRF — tính điểm và hợp nhất từ Dense + BM25 (k=44)", 14)

# ══════════════════════════════════════════════════════════════
# SECTION 1.5.x — after "1.5. Kỹ thuật đánh chỉ mục văn bản nâng cao"
# ══════════════════════════════════════════════════════════════
idx_15 = find_para_index(doc, "1.5. Kỹ thuật đánh chỉ mục")
end_15 = idx_15
for i in range(idx_15+1, len(doc.paragraphs)):
    p = doc.paragraphs[i]
    if p.style.name.startswith("Heading 2") and "1.6" in p.text:
        break
    end_15 = i

cur3 = end_15

# 1.5.1 HyDE
cur3 = insert_para_after(doc, cur3, "1.5.1. Biến đổi truy vấn — HyDE, Multi-Query và Decomposition", "Heading 3")
cur3 = insert_text_after(doc, cur3,
    "Trong thực tế, câu hỏi của người dùng thường ngắn và mang tính nghi vấn, trong khi chunk "
    "trong tài liệu dài và mang tính khẳng định. Khoảng cách ngữ nghĩa (semantic gap) này làm "
    "giảm chất lượng dense retrieval. Hệ thống áp dụng ba kỹ thuật biến đổi truy vấn để thu hẹp "
    "khoảng cách này:")
cur3 = insert_bullet_after(doc, cur3,
    "HyDE (Hypothetical Document Embeddings) [Gao et al., 2022]: LLM sinh ra một đoạn văn trả "
    "lời giả định (2-4 câu). Vector của đoạn giả định này gần với vector tài liệu thật hơn so "
    "với vector câu hỏi gốc, giúp dense search tìm đúng chunk hơn.")
cur3 = insert_bullet_after(doc, cur3,
    "Multi-Query Variants: LLM sinh tối đa 3 biến thể khác nhau của câu hỏi gốc (cùng ý định "
    "nhưng diễn đạt khác). Mỗi biến thể được embed và search độc lập, kết quả được gộp qua RRF. "
    "Tăng khả năng recall cho câu hỏi có thể diễn đạt nhiều cách.")
cur3 = insert_bullet_after(doc, cur3,
    "Query Decomposition: Với câu hỏi phức (so sánh, đa thực thể), LLM tách thành tối đa 2 "
    "sub-query đơn giản hơn. Mỗi sub-query được search riêng rồi tổng hợp lại. Giải quyết bài "
    "toán multi-hop reasoning.")
cur3 = insert_image_after(doc, cur3, "concept_hyde.png",
    "Hình 3.X+2. HyDE — thu hẹp khoảng cách ngữ nghĩa giữa câu hỏi và tài liệu", 14)

# 1.5.2 Reranking
cur3 = insert_para_after(doc, cur3, "1.5.2. Tái xếp hạng bằng LLM (LLM Reranking)", "Heading 3")
cur3 = insert_text_after(doc, cur3,
    "Sau khi RRF fusion tạo ra danh sách ứng viên, bước reranking sử dụng LLM để chấm điểm lại "
    "top-16 chunk theo mức độ liên quan thực sự với câu hỏi. Khác với bi-encoder (embedding model) "
    "xử lý câu hỏi và chunk riêng biệt, LLM reranking đọc đồng thời toàn bộ câu hỏi và nội dung "
    "chunk, cho phép nắm bắt các sắc thái ngữ nghĩa phức tạp mà vector search bỏ qua. "
    "Top-9 chunk có điểm cao nhất sau reranking được đưa vào context cho LLM sinh câu trả lời.")

# 1.5.3 OOS
cur3 = insert_para_after(doc, cur3, "1.5.3. Phát hiện câu hỏi ngoài phạm vi (OOS Detection)", "Heading 3")
cur3 = insert_text_after(doc, cur3,
    "Hệ thống triển khai cơ chế phát hiện câu hỏi ngoài phạm vi (Out-of-Scope) dựa trên điểm "
    "tương đồng vector của chunk được retrieve tốt nhất. Nếu điểm số cao nhất trong tập ứng viên "
    "thấp hơn ngưỡng OutOfScopeScoreThreshold=0.42, hệ thống từ chối trả lời thay vì sinh câu "
    "trả lời sai lệch. Ngoài ra, vùng điểm nằm giữa ngưỡng OOS và ClarifyScoreThreshold=0.61 "
    "kích hoạt cơ chế Clarify — hệ thống yêu cầu người dùng diễn đạt lại câu hỏi. "
    "Hai ngưỡng này được xác định qua Optuna và điều chỉnh thực nghiệm.")

# 1.5.4 Optuna
cur3 = insert_para_after(doc, cur3, "1.5.4. Tối ưu tham số tự động — Optuna TPE", "Heading 3")
cur3 = insert_text_after(doc, cur3,
    "Pipeline Document RAG có 11 tham số điều chỉnh ảnh hưởng đến chất lượng và tốc độ. Việc "
    "chỉnh tay từng tham số là bất khả thi do không gian tham số rộng và các tham số tương tác "
    "phi tuyến với nhau. Hệ thống sử dụng Optuna [Akiba et al., 2019] với thuật toán "
    "Tree-structured Parzen Estimator (TPE) để tự động tìm bộ tham số tối ưu.")
cur3 = insert_text_after(doc, cur3,
    "TPE hoạt động theo nguyên lý Bayesian optimization: từ 10 trial đầu random, thuật toán xây "
    "dựng hai phân phối xác suất l(x) (các tham số cho kết quả tốt) và g(x) (kết quả kém), sau "
    "đó đề xuất tham số mới từ vùng có tỉ lệ l(x)/g(x) cao. Objective function được thiết kế "
    "kết hợp 3 tiêu chí: accuracy (0.4), point recall (0.3), và refuse accuracy (0.2). "
    "Sau 50 trials, trial #22 đạt objective score 0.7321 với bộ tham số tối ưu được áp dụng "
    "vào hệ thống.")
cur3 = insert_image_after(doc, cur3, "concept_optuna_tpe.png",
    "Hình 3.X+3. Cơ chế TPE — xây dựng l(x)/g(x) và luồng 50 trials trên sweep2.db", 14)
cur3 = insert_image_after(doc, cur3, "optuna_convergence.png",
    "Hình 3.X+4. Đường hội tụ Optuna và fANOVA hyperparameter importance (50 trials)", 14)

# ══════════════════════════════════════════════════════════════
# SECTION 3.3.3–3.3.6 — after "3.3.2. Luồng xử lý chi tiết"
# ══════════════════════════════════════════════════════════════
idx_332 = find_para_index(doc, "3.3.2")
end_332 = idx_332
for i in range(idx_332+1, len(doc.paragraphs)):
    p = doc.paragraphs[i]
    if p.style.name.startswith("Heading") and p.text.strip() and "3.3" not in p.text:
        break
    end_332 = i

cur4 = end_332

# 3.3.3
cur4 = insert_para_after(doc, cur4, "3.3.3. Cài đặt giai đoạn Indexing Document RAG", "Heading 3")
cur4 = insert_text_after(doc, cur4,
    "Sau khi OCR trả về văn bản thô, hệ thống phân tích cấu trúc và nhận diện 3 loại nội dung: "
    "text thông thường, bảng biểu (table), và hình vẽ/sơ đồ (figure). Mỗi loại được đánh thẻ "
    "metadata content_type để phục vụ chiến lược tái xếp hạng theo loại nội dung (ContentType "
    "Reranking) ở bước sau.")
cur4 = insert_text_after(doc, cur4,
    "Văn bản được chunk theo chiến lược Parent-Child hai cấp: Parent chunk (~1600 ký tự, overlap "
    "240 ký tự) lưu trữ ngữ cảnh rộng; Child chunk (~700 ký tự, overlap 120 ký tự) phục vụ "
    "tìm kiếm chính xác. Mỗi chunk được mã hóa bằng mô hình multilingual-e5-small chạy ONNX "
    "Runtime trên CPU (384 chiều, mean pooling + L2 normalize) và lưu vào Qdrant Cloud collection "
    "document_vectors với Distance.Cosine. Hệ thống cũng duy trì collection rag_answer_cache "
    "để tái sử dụng câu trả lời cho câu hỏi tương tự (ngưỡng cosine ≥ 0.92).")
cur4 = insert_image_after(doc, cur4, "chart_indexing_structure.png",
    "Hình 3.X+5. Kiến trúc Parent-Child Chunking — từ OCR text đến Qdrant HNSW index", 14)

# 3.3.4
cur4 = insert_para_after(doc, cur4, "3.3.4. Cài đặt giai đoạn Retrieval — Multi-stage Pipeline", "Heading 3")
cur4 = insert_text_after(doc, cur4,
    "Khi nhận câu hỏi, hệ thống thực hiện 6 bước xử lý liên tiếp: (1) Query Transform — sinh "
    "multi-query variants, HyDE document, và sub-queries nếu cần; (2) Dense Search — Qdrant HNSW "
    "với ef_search tính động theo quy mô scope và số query (clamp 40–320); (3) BM25 Sparse Search "
    "— tự cài đặt TF saturation + IDF + length normalization; (4) RRF Fusion — hợp nhất Dense "
    "và BM25 score theo thứ hạng với k=44; (5) Cross-signal / Exact-signal / ContentType "
    "Reranking — điều chỉnh điểm theo tín hiệu phụ; (6) LLM Reranking — LLM chấm điểm top-16 "
    "và chọn top-9 chunk cuối cùng.")
cur4 = insert_image_after(doc, cur4, "chart_retrieval_pipeline.png",
    "Hình 3.X+6. Kiến trúc Multi-stage RAG Pipeline — toàn bộ luồng xử lý", 14)
cur4 = insert_image_after(doc, cur4, "chart_hnsw_ef_search.png",
    "Hình 3.X+7. Chiến lược ef_search động trong Qdrant HNSW theo scope và số query", 13)

# 3.3.5
cur4 = insert_para_after(doc, cur4, "3.3.5. Cơ chế Safety và UX", "Heading 3")
cur4 = insert_text_after(doc, cur4,
    "Hệ thống tích hợp bốn cơ chế an toàn và trải nghiệm người dùng: (1) OOS Detection — từ "
    "chối khi best score < 0.42, tránh sinh câu trả lời bịa đặt; (2) Clarify Detection — yêu "
    "cầu làm rõ khi score trong vùng 0.42–0.61, tránh trả lời mơ hồ; (3) Ambiguity Detection "
    "— khi câu hỏi khớp với nhiều tài liệu khác nhau trong workspace, hệ thống yêu cầu người "
    "dùng chọn cụ thể tài liệu nào; (4) Conversation History RAG — lưu 3 lượt hội thoại gần "
    "nhất dưới dạng vector trong collection conversation_history, cho phép hỏi đáp liên tục "
    "có ngữ cảnh mà không cần nhắc lại.")

# 3.3.6
cur4 = insert_para_after(doc, cur4, "3.3.6. Tối ưu tham số — Optuna Sweep", "Heading 3")
cur4 = insert_text_after(doc, cur4,
    "11 tham số của RAG Pipeline được tối ưu tự động qua Optuna TPE với 50 trials lưu trong "
    "sweep2.db. Objective function: 0.4 × accuracy + 0.3 × point_recall + 0.2 × refuse_accuracy. "
    "Trial #22 đạt objective tốt nhất 0.7321 với bộ tham số: RetrievePerQuery=23, "
    "CandidatePoolLimit=30, RerankCandidateLimit=16, QueryVariantLimit=3, "
    "DecompositionSubQueryLimit=2, RrfK=44, Bm25K1=1.343, Bm25B=0.574, TopK=9. "
    "Tham số OutOfScopeScoreThreshold được điều chỉnh thủ công từ 0.539 (Optuna) xuống 0.42 "
    "dựa trên quan sát thực tế về trade-off giữa false refusal và OOS detection.")
cur4 = insert_image_after(doc, cur4, "chart_params_range.png",
    "Hình 3.X+8. Vị trí 11 tham số tối ưu trong không gian tìm kiếm Optuna", 13)

# ─── Save ───────────────────────────────────────────────────
doc.save("baocao.docx")
print("Saved baocao.docx successfully!")
print(f"Total paragraphs: {len(doc.paragraphs)}")
